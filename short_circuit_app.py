#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Short Circuit Calculator Application
Генератор курсовой работы по расчёту токов короткого замыкания
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import math
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from PIL import Image
import os

class ShortCircuitCalculator:
    def __init__(self, root):
        self.root = root
        self.root.title('Расчёт токов короткого замыкания - Курсовая работа')
        self.root.geometry('800x600')
        
        # Переменные для хранения данных
        self.data = {}
        
        self.create_widgets()
        
    def create_widgets(self):
        # Создаем notebook для вкладок
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Вкладка 1: Исходные данные
        self.input_frame = ttk.Frame(notebook)
        notebook.add(self.input_frame, text='Исходные данные')
        self.create_input_tab()
        
        # Вкладка 2: Результаты расчёта
        self.results_frame = ttk.Frame(notebook)
        notebook.add(self.results_frame, text='Результаты расчёта')
        self.create_results_tab()
        
        # Вкладка 3: Графики
        self.graphs_frame = ttk.Frame(notebook)
        notebook.add(self.graphs_frame, text='Графики и диаграммы')
        self.create_graphs_tab()
        
        # Кнопка генерации DOCX
        btn_generate = tk.Button(self.root, text='Генерировать DOCX файл', 
                               command=self.generate_docx, bg='#2c3e50', fg='white',
                               font=('Arial', 12, 'bold'))
        btn_generate.pack(pady=10)
        
    def create_input_tab(self):
        # Создаем фрейм с прокруткой
        canvas = tk.Canvas(self.input_frame)
        scrollbar = ttk.Scrollbar(self.input_frame, orient='vertical', command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            '<Configure>',
            lambda e: canvas.configure(scrollregion=canvas.bbox('all'))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Параметры генератора
        gen_frame = ttk.LabelFrame(scrollable_frame, text='Параметры генератора', padding=10)
        gen_frame.grid(row=0, column=0, columnspan=2, sticky='ew', pady=5, padx=10)
        
        tk.Label(gen_frame, text='Напряжение (кВ):').grid(row=0, column=0, sticky='w', pady=2)
        self.voltage_var = tk.StringVar(value='110')
        tk.Entry(gen_frame, textvariable=self.voltage_var, width=15).grid(row=0, column=1, sticky='w', pady=2)
        
        tk.Label(gen_frame, text='Мощность (МВт):').grid(row=1, column=0, sticky='w', pady=2)
        self.power_var = tk.StringVar(value='20')
        tk.Entry(gen_frame, textvariable=self.power_var, width=15).grid(row=1, column=1, sticky='w', pady=2)
        
        tk.Label(gen_frame, text='Сверхпереходное сопротивление (о.е.):').grid(row=2, column=0, sticky='w', pady=2)
        self.xd_var = tk.StringVar(value='0.22')
        tk.Entry(gen_frame, textvariable=self.xd_var, width=15).grid(row=2, column=1, sticky='w', pady=2)
        
        tk.Label(gen_frame, text='Коэффициент мощности:').grid(row=3, column=0, sticky='w', pady=2)
        self.cosphi_var = tk.StringVar(value='0.8')
        tk.Entry(gen_frame, textvariable=self.cosphi_var, width=15).grid(row=3, column=1, sticky='w', pady=2)
        
        # Параметры трансформатора
        trans_frame = ttk.LabelFrame(scrollable_frame, text='Параметры трансформатора', padding=10)
        trans_frame.grid(row=1, column=0, columnspan=2, sticky='ew', pady=5, padx=10)
        
        tk.Label(trans_frame, text='Мощность (МВА):').grid(row=0, column=0, sticky='w', pady=2)
        self.trans_power_var = tk.StringVar(value='25')
        tk.Entry(trans_frame, textvariable=self.trans_power_var, width=15).grid(row=0, column=1, sticky='w', pady=2)
        
        tk.Label(trans_frame, text='Напряжение КЗ (%):').grid(row=1, column=0, sticky='w', pady=2)
        self.uk_var = tk.StringVar(value='10.5')
        tk.Entry(trans_frame, textvariable=self.uk_var, width=15).grid(row=1, column=1, sticky='w', pady=2)
        
        # Параметры линии
        line_frame = ttk.LabelFrame(scrollable_frame, text='Параметры линии', padding=10)
        line_frame.grid(row=2, column=0, columnspan=2, sticky='ew', pady=5, padx=10)
        
        tk.Label(line_frame, text='Длина (км):').grid(row=0, column=0, sticky='w', pady=2)
        self.line_length_var = tk.StringVar(value='61')
        tk.Entry(line_frame, textvariable=self.line_length_var, width=15).grid(row=0, column=1, sticky='w', pady=2)
        
        tk.Label(line_frame, text='Удельное сопротивление (Ом/км):').grid(row=1, column=0, sticky='w', pady=2)
        self.line_res_var = tk.StringVar(value='0.429')
        tk.Entry(line_frame, textvariable=self.line_res_var, width=15).grid(row=1, column=1, sticky='w', pady=2)
        
        # Базисные величины
        base_frame = ttk.LabelFrame(scrollable_frame, text='Базисные величины', padding=10)
        base_frame.grid(row=3, column=0, columnspan=2, sticky='ew', pady=5, padx=10)
        
        tk.Label(base_frame, text='Базисная мощность (МВА):').grid(row=0, column=0, sticky='w', pady=2)
        self.base_power_var = tk.StringVar(value='1000')
        tk.Entry(base_frame, textvariable=self.base_power_var, width=15).grid(row=0, column=1, sticky='w', pady=2)
        
        tk.Label(base_frame, text='Базисное напряжение (кВ):').grid(row=1, column=0, sticky='w', pady=2)
        self.base_voltage_var = tk.StringVar(value='110')
        tk.Entry(base_frame, textvariable=self.base_voltage_var, width=15).grid(row=1, column=1, sticky='w', pady=2)
        
        # Кнопка расчёта
        btn_calc = tk.Button(scrollable_frame, text='Выполнить расчёт', 
                           command=self.calculate, bg='#27ae60', fg='white',
                           font=('Arial', 10, 'bold'))
        btn_calc.grid(row=4, column=0, columnspan=2, pady=20)
        
        canvas.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
    def create_results_tab(self):
        # Текстовое поле для результатов
        self.results_text = tk.Text(self.results_frame, wrap='word', font=('Courier', 10))
        scrollbar_results = ttk.Scrollbar(self.results_frame, orient='vertical', command=self.results_text.yview)
        self.results_text.configure(yscrollcommand=scrollbar_results.set)
        
        self.results_text.pack(side='left', fill='both', expand=True, padx=10, pady=10)
        scrollbar_results.pack(side='right', fill='y')
        
    def create_graphs_tab(self):
        # Фрейм для графиков
        self.graphs_canvas = tk.Canvas(self.graphs_frame, bg='white')
        self.graphs_canvas.pack(fill='both', expand=True, padx=10, pady=10)
        
    def calculate(self):
        try:
            # Получаем данные из полей
            voltage = float(self.voltage_var.get())
            power = float(self.power_var.get())
            xd = float(self.xd_var.get())
            cosphi = float(self.cosphi_var.get())
            trans_power = float(self.trans_power_var.get())
            uk = float(self.uk_var.get())
            line_length = float(self.line_length_var.get())
            line_res = float(self.line_res_var.get())
            base_power = float(self.base_power_var.get())
            base_voltage = float(self.base_voltage_var.get())
            
            # Выполняем расчёты
            results = self.perform_calculations(voltage, power, xd, cosphi, trans_power, uk, 
                                              line_length, line_res, base_power, base_voltage)
            
            # Сохраняем результаты
            self.data = results
            
            # Отображаем результаты
            self.display_results(results)
            
            # Создаём графики
            self.create_graphs(results)
            
            messagebox.showinfo('Успех', 'Расчёт выполнен успешно!')
            
        except Exception as e:
            messagebox.showerror('Ошибка', f'Ошибка при расчёте: {str(e)}')
            
    def perform_calculations(self, voltage, power, xd, cosphi, trans_power, uk, 
                           line_length, line_res, base_power, base_voltage):
        # Расчёт сверхпереходной ЭДС генератора
        sinphi = math.sqrt(1 - cosphi**2)
        generator_emf = math.sqrt(cosphi**2 + (sinphi + xd)**2) * voltage
        
        # Расчёт сопротивлений
        generator_resistance = xd * (voltage**2 / power) * (base_power / voltage**2)
        transformer_resistance = (uk / 100) * (voltage**2 / trans_power)
        line_resistance = line_res * line_length
        
        total_resistance = generator_resistance + transformer_resistance + line_resistance
        
        # Расчёт токов короткого замыкания
        three_phase_current = generator_emf / (math.sqrt(3) * total_resistance)
        two_phase_current = (math.sqrt(3) / 2) * three_phase_current
        single_phase_current = 3 * generator_emf / (3 * total_resistance)
        two_phase_ground_current = math.sqrt(3) * generator_emf / (2 * total_resistance)
        
        # Симметричные составляющие
        positive_sequence_current = three_phase_current
        negative_sequence_current = -three_phase_current
        zero_sequence_current = 0
        
        # Напряжения в месте КЗ
        positive_sequence_voltage = generator_emf - positive_sequence_current * total_resistance
        negative_sequence_voltage = -negative_sequence_current * total_resistance
        zero_sequence_voltage = -zero_sequence_current * total_resistance
        
        # Коэффициенты токораспределения
        generator_coefficient = generator_resistance / total_resistance
        transformer_coefficient = transformer_resistance / total_resistance
        line_coefficient = line_resistance / total_resistance
        
        return {
            'voltage': voltage,
            'power': power,
            'xd': xd,
            'cosphi': cosphi,
            'trans_power': trans_power,
            'uk': uk,
            'line_length': line_length,
            'line_res': line_res,
            'base_power': base_power,
            'base_voltage': base_voltage,
            'generator_emf': generator_emf,
            'generator_resistance': generator_resistance,
            'transformer_resistance': transformer_resistance,
            'line_resistance': line_resistance,
            'total_resistance': total_resistance,
            'three_phase_current': three_phase_current,
            'two_phase_current': two_phase_current,
            'single_phase_current': single_phase_current,
            'two_phase_ground_current': two_phase_ground_current,
            'positive_sequence_current': positive_sequence_current,
            'negative_sequence_current': negative_sequence_current,
            'zero_sequence_current': zero_sequence_current,
            'positive_sequence_voltage': positive_sequence_voltage,
            'negative_sequence_voltage': negative_sequence_voltage,
            'zero_sequence_voltage': zero_sequence_voltage,
            'generator_coefficient': generator_coefficient,
            'transformer_coefficient': transformer_coefficient,
            'line_coefficient': line_coefficient
        }
        
    def display_results(self, results):
        self.results_text.delete(1.0, tk.END)
        
        output = f'''
РАСЧЁТ ТОКОВ КОРОТКОГО ЗАМЫКАНИЯ В ЭЛЕКТРОЭНЕРГЕТИЧЕСКОЙ СИСТЕМЕ
================================================================

ИСХОДНЫЕ ДАННЫЕ:
----------------
Напряжение системы: {results['voltage']} кВ
Мощность генератора: {results['power']} МВт
Сверхпереходное сопротивление генератора: {results['xd']} о.е.
Коэффициент мощности: {results['cosphi']}
Мощность трансформатора: {results['trans_power']} МВА
Напряжение короткого замыкания трансформатора: {results['uk']} %
Длина линии: {results['line_length']} км
Удельное сопротивление линии: {results['line_res']} Ом/км
Базисная мощность: {results['base_power']} МВА
Базисное напряжение: {results['base_voltage']} кВ

ПАРАМЕТРЫ СХЕМЫ ЗАМЕЩЕНИЯ:
--------------------------
Сверхпереходная ЭДС генератора: {results['generator_emf']:.2f} кВ
Сопротивление генератора: {results['generator_resistance']:.3f} Ом
Сопротивление трансформатора: {results['transformer_resistance']:.3f} Ом
Сопротивление линии: {results['line_resistance']:.3f} Ом
Общее сопротивление: {results['total_resistance']:.3f} Ом

ТОКИ КОРОТКОГО ЗАМЫКАНИЯ:
-------------------------
Ток трёхфазного КЗ: {results['three_phase_current']:.2f} кА
Ток двухфазного КЗ: {results['two_phase_current']:.2f} кА
Ток однофазного КЗ: {results['single_phase_current']:.2f} кА
Ток двухфазного КЗ на землю: {results['two_phase_ground_current']:.2f} кА

СИММЕТРИЧНЫЕ СОСТАВЛЯЮЩИЕ ТОКОВ:
--------------------------------
Ток прямой последовательности: {results['positive_sequence_current']:.2f} кА
Ток обратной последовательности: {results['negative_sequence_current']:.2f} кА
Ток нулевой последовательности: {results['zero_sequence_current']:.2f} кА

СИММЕТРИЧНЫЕ СОСТАВЛЯЮЩИЕ НАПРЯЖЕНИЙ:
------------------------------------
Напряжение прямой последовательности: {results['positive_sequence_voltage']:.2f} кВ
Напряжение обратной последовательности: {results['negative_sequence_voltage']:.2f} кВ
Напряжение нулевой последовательности: {results['zero_sequence_voltage']:.2f} кВ

КОЭФФИЦИЕНТЫ ТОКОРАСПРЕДЕЛЕНИЯ:
-------------------------------
Коэффициент генератора: {results['generator_coefficient']:.3f}
Коэффициент трансформатора: {results['transformer_coefficient']:.3f}
Коэффициент линии: {results['line_coefficient']:.3f}
'''
        
        self.results_text.insert(1.0, output)
        
    def create_graphs(self, results):
        # Очищаем canvas
        self.graphs_canvas.delete('all')
        
        # Создаём график токов КЗ
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
        
        # График 1: Токи КЗ
        current_types = ['3-фазное', '2-фазное', '1-фазное', '2-фазное на землю']
        current_values = [
            results['three_phase_current'],
            results['two_phase_current'], 
            results['single_phase_current'],
            results['two_phase_ground_current']
        ]
        
        bars1 = ax1.bar(current_types, current_values, color=['#e74c3c', '#f39c12', '#27ae60', '#3498db'])
        ax1.set_title('Токи короткого замыкания', fontsize=14, fontweight='bold')
        ax1.set_ylabel('Ток (кА)', fontsize=12)
        ax1.grid(True, alpha=0.3)
        
        # Добавляем значения на столбцы
        for bar, value in zip(bars1, current_values):
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                    f'{value:.2f} кА', ha='center', va='bottom', fontweight='bold')
        
        # График 2: Симметричные составляющие
        sequence_types = ['Прямая', 'Обратная', 'Нулевая']
        sequence_currents = [
            results['positive_sequence_current'],
            results['negative_sequence_current'],
            results['zero_sequence_current']
        ]
        
        bars2 = ax2.bar(sequence_types, sequence_currents, color=['#e74c3c', '#f39c12', '#27ae60'])
        ax2.set_title('Симметричные составляющие токов', fontsize=14, fontweight='bold')
        ax2.set_ylabel('Ток (кА)', fontsize=12)
        ax2.grid(True, alpha=0.3)
        
        # Добавляем значения на столбцы
        for bar, value in zip(bars2, sequence_currents):
            height = bar.get_height()
            ax2.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                    f'{value:.2f} кА', ha='center', va='bottom', fontweight='bold')
        
        plt.tight_layout()
        
        # Сохраняем график
        plt.savefig('short_circuit_graphs.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # Отображаем график в canvas
        img = Image.open('short_circuit_graphs.png')
        img = img.resize((700, 400), Image.Resampling.LANCZOS)
        
        # Конвертируем в PhotoImage для tkinter
        import tkinter as tk
        self.graph_image = tk.PhotoImage(img)
        self.graphs_canvas.create_image(350, 200, image=self.graph_image)
        
    def generate_docx(self):
        if not self.data:
            messagebox.showerror('Ошибка', 'Сначала выполните расчёт!')
            return
            
        try:
            # Создаём документ
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('РАСЧЁТ ТОКОВ КОРОТКОГО ЗАМЫКАНИЯ В ЭЛЕКТРОЭНЕРГЕТИЧЕСКОЙ СИСТЕМЕ', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Информация о курсовой работе
            doc.add_heading('Курсовая работа по дисциплине "Переходные процессы"', level=1)
            doc.add_paragraph('Направление подготовки: 13.03.02 Электроэнергетика и электротехника')
            doc.add_paragraph('Выполнил: студент 4 курса, группы 13-НФ21 Корытников Р.М')
            doc.add_paragraph('Руководитель: к.т.н., доцент по кафедре "ЭЭиАТП" Складчиков А.А.')
            
            # Исходные данные
            doc.add_heading('1. ИСХОДНЫЕ ДАННЫЕ ДЛЯ ПРОЕКТИРОВАНИЯ', level=1)
            
            # Таблица исходных данных
            table1 = doc.add_table(rows=1, cols=2)
            table1.style = 'Table Grid'
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table1.rows[0].cells
            hdr_cells[0].text = 'Параметр'
            hdr_cells[1].text = 'Значение'
            
            data_rows = [
                ('Напряжение системы', f'{self.data["voltage"]} кВ'),
                ('Мощность генератора', f'{self.data["power"]} МВт'),
                ('Сверхпереходное сопротивление генератора', f'{self.data["xd"]} о.е.'),
                ('Коэффициент мощности', f'{self.data["cosphi"]}'),
                ('Мощность трансформатора', f'{self.data["trans_power"]} МВА'),
                ('Напряжение короткого замыкания трансформатора', f'{self.data["uk"]} %'),
                ('Длина линии', f'{self.data["line_length"]} км'),
                ('Удельное сопротивление линии', f'{self.data["line_res"]} Ом/км'),
                ('Базисная мощность', f'{self.data["base_power"]} МВА'),
                ('Базисное напряжение', f'{self.data["base_voltage"]} кВ')
            ]
            
            for param, value in data_rows:
                row_cells = table1.add_row().cells
                row_cells[0].text = param
                row_cells[1].text = value
            
            # Расчёт параметров схемы замещения
            doc.add_heading('2. РАСЧЁТ ПАРАМЕТРОВ СХЕМЫ ЗАМЕЩЕНИЯ', level=1)
            
            doc.add_paragraph('Сверхпереходная ЭДС генератора рассчитывается по формуле:')
            doc.add_paragraph('E\'\' = √(cos²φ + (sinφ + xd\'\')²) × Uном')
            
            doc.add_paragraph(f'где: cosφ = {self.data["cosphi"]}, xd\'\' = {self.data["xd"]} о.е., Uном = {self.data["voltage"]} кВ')
            
            sinphi = math.sqrt(1 - self.data['cosphi']**2)
            doc.add_paragraph(f'sinφ = √(1 - cos²φ) = √(1 - {self.data["cosphi"]}²) = {sinphi:.3f}')
            
            doc.add_paragraph(f'E\'\' = √({self.data["cosphi"]}² + ({sinphi:.3f} + {self.data["xd"]})²) × {self.data["voltage"]} = {self.data["generator_emf"]:.2f} кВ')
            
            # Таблица параметров схемы замещения
            table2 = doc.add_table(rows=1, cols=2)
            table2.style = 'Table Grid'
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells2 = table2.rows[0].cells
            hdr_cells2[0].text = 'Параметр'
            hdr_cells2[1].text = 'Значение'
            
            scheme_data = [
                ('Сверхпереходная ЭДС генератора', f'{self.data["generator_emf"]:.2f} кВ'),
                ('Сопротивление генератора', f'{self.data["generator_resistance"]:.3f} Ом'),
                ('Сопротивление трансформатора', f'{self.data["transformer_resistance"]:.3f} Ом'),
                ('Сопротивление линии', f'{self.data["line_resistance"]:.3f} Ом'),
                ('Общее сопротивление', f'{self.data["total_resistance"]:.3f} Ом')
            ]
            
            for param, value in scheme_data:
                row_cells = table2.add_row().cells
                row_cells[0].text = param
                row_cells[1].text = value
            
            # Расчёт токов короткого замыкания
            doc.add_heading('3. РАСЧЁТ ТОКОВ КОРОТКОГО ЗАМЫКАНИЯ', level=1)
            
            doc.add_paragraph('Ток трёхфазного короткого замыкания:')
            doc.add_paragraph('Iкз(3) = E\'\' / (√3 × Z)')
            doc.add_paragraph(f'Iкз(3) = {self.data["generator_emf"]:.2f} / (√3 × {self.data["total_resistance"]:.3f}) = {self.data["three_phase_current"]:.2f} кА')
            
            doc.add_paragraph('Ток двухфазного короткого замыкания:')
            doc.add_paragraph('Iкз(2) = (√3/2) × Iкз(3)')
            doc.add_paragraph(f'Iкз(2) = (√3/2) × {self.data["three_phase_current"]:.2f} = {self.data["two_phase_current"]:.2f} кА')
            
            # Таблица токов КЗ
            table3 = doc.add_table(rows=1, cols=2)
            table3.style = 'Table Grid'
            table3.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells3 = table3.rows[0].cells
            hdr_cells3[0].text = 'Вид КЗ'
            hdr_cells3[1].text = 'Ток (кА)'
            
            current_data = [
                ('Трёхфазное КЗ', f'{self.data["three_phase_current"]:.2f}'),
                ('Двухфазное КЗ', f'{self.data["two_phase_current"]:.2f}'),
                ('Однофазное КЗ', f'{self.data["single_phase_current"]:.2f}'),
                ('Двухфазное КЗ на землю', f'{self.data["two_phase_ground_current"]:.2f}')
            ]
            
            for kz_type, current in current_data:
                row_cells = table3.add_row().cells
                row_cells[0].text = kz_type
                row_cells[1].text = current
            
            # Симметричные составляющие
            doc.add_heading('4. СИММЕТРИЧНЫЕ СОСТАВЛЯЮЩИЕ', level=1)
            
            # Таблица симметричных составляющих
            table4 = doc.add_table(rows=1, cols=4)
            table4.style = 'Table Grid'
            table4.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells4 = table4.rows[0].cells
            hdr_cells4[0].text = 'Последовательность'
            hdr_cells4[1].text = 'Ток (кА)'
            hdr_cells4[2].text = 'Напряжение (кВ)'
            hdr_cells4[3].text = 'Коэффициент'
            
            sequence_data = [
                ('Прямая', f'{self.data["positive_sequence_current"]:.2f}', f'{self.data["positive_sequence_voltage"]:.2f}', f'{self.data["generator_coefficient"]:.3f}'),
                ('Обратная', f'{self.data["negative_sequence_current"]:.2f}', f'{self.data["negative_sequence_voltage"]:.2f}', f'{self.data["transformer_coefficient"]:.3f}'),
                ('Нулевая', f'{self.data["zero_sequence_current"]:.2f}', f'{self.data["zero_sequence_voltage"]:.2f}', f'{self.data["line_coefficient"]:.3f}')
            ]
            
            for seq, current, voltage, coeff in sequence_data:
                row_cells = table4.add_row().cells
                row_cells[0].text = seq
                row_cells[1].text = current
                row_cells[2].text = voltage
                row_cells[3].text = coeff
            
            # Добавляем график
            doc.add_heading('5. ГРАФИКИ И ДИАГРАММЫ', level=1)
            
            if os.path.exists('short_circuit_graphs.png'):
                doc.add_picture('short_circuit_graphs.png', width=Inches(6))
            
            # Заключение
            doc.add_heading('6. ЗАКЛЮЧЕНИЕ', level=1)
            
            max_current = max(self.data['three_phase_current'], self.data['two_phase_current'], 
                            self.data['single_phase_current'], self.data['two_phase_ground_current'])
            
            max_current_type = 'трёхфазном' if self.data['three_phase_current'] == max_current else \
                              'двухфазном' if self.data['two_phase_current'] == max_current else \
                              'однофазном' if self.data['single_phase_current'] == max_current else 'двухфазном на землю'
            
            max_resistance = max(self.data['generator_coefficient'], self.data['transformer_coefficient'], 
                               self.data['line_coefficient'])
            
            max_resistance_type = 'генератор' if self.data['generator_coefficient'] == max_resistance else \
                                 'трансформатор' if self.data['transformer_coefficient'] == max_resistance else 'линия'
            
            conclusion = f'''
По результатам расчётов токов короткого замыкания в электроэнергетической системе получены следующие основные результаты:

1. Максимальный ток короткого замыкания составляет {max_current:.2f} кА при {max_current_type} КЗ.

2. Сопротивление генератора составляет {self.data["generator_resistance"]:.3f} Ом, что соответствует {self.data["generator_coefficient"]*100:.1f}% от общего сопротивления системы.

3. Наибольший вклад в общее сопротивление вносят: {max_resistance_type} ({max_resistance*100:.1f}%).

4. Полученные значения токов КЗ позволяют правильно выбрать и настроить защитное оборудование для обеспечения надёжной работы энергосистемы.
'''
            
            doc.add_paragraph(conclusion)
            
            # Сохраняем документ
            filename = filedialog.asksaveasfilename(
                defaultextension='.docx',
                filetypes=[('Word документы', '*.docx'), ('Все файлы', '*.*')],
                title='Сохранить расчёт как...'
            )
            
            if filename:
                doc.save(filename)
                messagebox.showinfo('Успех', f'DOCX файл сохранён: {filename}')
                
        except Exception as e:
            messagebox.showerror('Ошибка', f'Ошибка при создании DOCX файла: {str(e)}')

if __name__ == '__main__':
    root = tk.Tk()
    app = ShortCircuitCalculator(root)
    root.mainloop()


